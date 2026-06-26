"""Build 报告.docx from doc/报告.md with the course-required formatting.

Requirements (doc/Paper_Format.md §1.2): 宋体 小四 (12pt) body, 20pt EXACT line
spacing, first-line indent 2 characters (w:firstLineChars=200), ≥3000 Chinese
characters, figures/tables encouraged. Headings use 黑体 (SimHei) bold; tables use
a three-line style; images are centered with a caption.

This is a deliberately small Markdown subset parser tuned to our own report.md:
title (#), section (##), subsection (###), paragraphs, **bold** inline,
| pipe tables |, ![caption](path) images, ``` fenced pseudocode, > notes,
- bullets, and [n] references.

Usage:  python scripts/build_report_docx.py [doc/报告.md] [报告.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

REPO = Path("/inspire/hdd/project/video-understanding/public/personal/chwang/live/cpjs/BD_Hadoop")
BODY = "宋体"; HEAD = "黑体"; MONO = "Consolas"; WEST = "Times New Roman"


def set_run_font(run, size, bold=False, east=BODY, west=WEST, color=None):
    run.font.size = Pt(size); run.bold = bold
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), west); rf.set(qn("w:hAnsi"), west); rf.set(qn("w:eastAsia"), east)
    if color is not None:
        run.font.color.rgb = color


LINK = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
CITE = re.compile(r"(\[\d+\])")               # inline reference marker, e.g. [12]


def _plain(p, seg, size, bold, east, cite):
    """Emit a non-code text run; if cite, raise [n] markers to superscript."""
    if not cite:
        set_run_font(p.add_run(seg), size, bold=bold, east=east); return
    for part in CITE.split(seg):
        if part == "":
            continue
        if CITE.fullmatch(part):
            r = p.add_run(part); set_run_font(r, size - 3, bold=bold, east=east)
            r.font.superscript = True                         # citation as right-top superscript
        else:
            set_run_font(p.add_run(part), size, bold=bold, east=east)


def _emit(p, text, size, east, cite=False):
    """Render inline **bold**, `code`, and (optionally) [n] superscript citations."""
    for bi, bold_seg in enumerate(text.split("**")):
        if bold_seg == "":
            continue
        bold = (bi % 2 == 1)
        for ci, seg in enumerate(bold_seg.split("`")):
            if seg == "":
                continue
            if ci % 2 == 1:                                   # inside backticks
                set_run_font(p.add_run(seg), size - 0.5, bold=bold, east=MONO, west=MONO)
            else:
                _plain(p, seg, size, bold, east, cite)


def add_hyperlink(p, url, text, size):
    rid = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    for tag, attr, val in [("w:color", "w:val", "0563C1"), ("w:u", "w:val", "single")]:
        e = OxmlElement(tag); e.set(qn(attr), val); rpr.append(e)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), WEST); rf.set(qn("w:hAnsi"), WEST); rf.set(qn("w:eastAsia"), BODY); rpr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hl.append(run); p._p.append(hl)


def add_runs_bold(p, text, size, east=BODY, cite=False):
    """Render inline [text](url) links, **bold**, `code`, and [n] citations."""
    pos = 0
    for m in LINK.finditer(text):
        _emit(p, text[pos:m.start()], size, east, cite)
        add_hyperlink(p, m.group(2), m.group(1), size)
        pos = m.end()
    _emit(p, text[pos:], size, east, cite)


def fmt_body(p, indent=True):
    pf = p.paragraph_format
    pf.line_spacing = Pt(20); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Pt(24)                       # 2 chars × 12pt fallback
        ind = p._p.get_or_add_pPr().find(qn("w:ind"))
        if ind is not None:
            ind.set(qn("w:firstLineChars"), "200")          # exact "2 字符"


def border(el, edge, sz="6", color="000000"):
    tcb = el.find(qn("w:tcBorders"))
    if tcb is None:
        tcb = OxmlElement("w:tcBorders")
        el.append(tcb)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), sz)
    e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
    tcb.append(e)


def shade(tcPr, fill="F5F6F8"):
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)


def cell_margins(cell, top=80, bottom=80, left=140, right=120):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement("w:tcMar")
    for tag, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); m.append(e)
    tcPr.append(m)


def fixed_widths(t, widths):
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    for row in t.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths):
                cell.width = widths[j]


def three_line(table):
    rows = table.rows
    for cell in rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr(); border(tcPr, "top", "12"); border(tcPr, "bottom", "6")
    for cell in rows[-1].cells:
        tcPr = cell._tc.get_or_add_tcPr(); border(tcPr, "bottom", "12")


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows
             if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(18); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            add_runs_bold(p, val if "**" in val else (f"**{val}**" if i == 0 else val), 10.5)
    if "签名" in cells[0]:                                   # member table → leave room to sign
        fixed_widths(t, [Cm(2.0), Cm(2.6), Cm(8.4), Cm(2.4)])
    three_line(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, caption, path, width=11.5):
    fp = (REPO / path) if not Path(path).is_absolute() else Path(path)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6); pf.space_after = Pt(2); pf.keep_with_next = True   # stay with caption
    if fp.exists():
        p.add_run().add_picture(str(fp), width=Cm(width))
    else:
        set_run_font(p.add_run(f"[缺图: {path}]"), 10)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(1); cap.paragraph_format.space_after = Pt(8)
    set_run_font(cap.add_run(caption), 9)


def add_grid(doc, caption, items, page_break=True, img_w=7.4):
    """items: list of (path, subcaption); lay out 2 per row, borderless, then caption.
    page_break=True puts the block on a fresh page (the 6-panel board); False keeps
    it inline (a 2-up row to pack figures and cut whitespace)."""
    if page_break:
        doc.add_page_break()
    n = len(items); cols = 2; rows = (n + cols - 1) // cols
    t = doc.add_table(rows=rows * 2, cols=cols); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:                                   # keep each row off page breaks
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit"); trPr.append(cant)
    for idx, (path, sub) in enumerate(items):
        r, c = (idx // cols) * 2, idx % cols
        fp = (REPO / path) if not Path(path).is_absolute() else Path(path)
        cell = t.cell(r, c); cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(0)
        if fp.exists():
            cp.add_run().add_picture(str(fp), width=Cm(img_w))
        scell = t.cell(r + 1, c); sp = scell.paragraphs[0]; sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(6)
        set_run_font(sp.add_run(sub), 8.5)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2); cap.paragraph_format.space_after = Pt(10)
    set_run_font(cap.add_run(caption), 9, bold=True)


def add_stack(doc, caption, paths, img_w=11.0):
    """Stack images in one tight column (no inter-image whitespace), one caption."""
    t = doc.add_table(rows=len(paths), cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, path in enumerate(paths):
        trPr = t.rows[k]._tr.get_or_add_trPr(); trPr.append(OxmlElement("w:cantSplit"))
        cell = t.cell(k, 0); cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(1); cp.paragraph_format.space_after = Pt(1)
        fp = (REPO / path) if not Path(path).is_absolute() else Path(path)
        if fp.exists():
            cp.add_run().add_picture(str(fp), width=Cm(img_w))
        else:
            set_run_font(cp.add_run(f"[缺图: {path}]"), 10)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2); cap.paragraph_format.space_after = Pt(8)
    set_run_font(cap.add_run(caption), 9, bold=True)


def add_code(doc, lines):
    """Render a pseudocode block as a shaded, bordered, roomy code box."""
    spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.line_spacing = Pt(4); spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0); tcPr = cell._tc.get_or_add_tcPr()
    shade(tcPr, "F5F6F8")
    for e in ("top", "bottom", "left", "right"):
        border(tcPr, e, "4", "C8CDD7")
    cell_margins(cell, top=90, bottom=90, left=150, right=120)
    for k, ln in enumerate(lines):
        p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = Pt(15.5); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        set_run_font(p.add_run(ln if ln else " "), 10.5, east=MONO, west=MONO)
    after = doc.add_paragraph(); after.paragraph_format.space_after = Pt(6)
    after.paragraph_format.line_spacing = Pt(4); after.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def heading(doc, text, level):
    sizes = {1: 15, 2: 13.5}
    before = {0: 6, 1: 20, 2: 14}        # more breathing room between sections
    after = {0: 10, 1: 9, 2: 6}
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(before.get(level, 14)); pf.space_after = Pt(after.get(level, 6))
    pf.line_spacing = Pt(24); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), sizes.get(level, 18) if level else 18, bold=True, east=HEAD)


def build(md_path, out_path):
    doc = Document()
    # page-wide default + margins
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.8); sec.top_margin = sec.bottom_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = WEST; style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)

    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    cjk = 0
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip("\n")
        s = ln.strip()
        if not s:
            i += 1; continue
        if s.startswith("```"):                                  # code fence
            block = []; i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i]); i += 1
            add_code(doc, block); i += 1; continue
        if s.startswith("@@grid|") or s.startswith("@@row|"):     # image grid / inline 2-up row
            is_row = s.startswith("@@row|")
            cap = s.split("|", 1)[1]; items = []; i += 1
            while i < len(lines) and lines[i].strip() != "@@":
                parts = lines[i].strip().split("|")
                if len(parts) == 2:
                    items.append((parts[0], parts[1]))
                i += 1
            add_grid(doc, cap, items, page_break=not is_row, img_w=7.7 if is_row else 7.4)
            i += 1; continue
        if s.startswith("@@stack|"):                             # tight 1-column image stack
            cap = s.split("|", 1)[1]; paths = []; i += 1
            while i < len(lines) and lines[i].strip() != "@@":
                if lines[i].strip():
                    paths.append(lines[i].strip())
                i += 1
            add_stack(doc, cap, paths); i += 1; continue
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", s)                # image
        if m:
            p2 = m.group(2)
            w = 10.0 if ("map" in p2 or "flow" in p2) else (10.5 if "demo_shots" in p2 else 11.5)
            add_image(doc, m.group(1), p2, w); i += 1; continue
        if s.startswith("|"):                                    # table block
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip()); i += 1
            add_table(doc, block); continue
        if s.startswith("# "):
            heading(doc, s[2:].strip(), 0); cjk += count_cjk(s); i += 1; continue
        if s.startswith("## "):
            heading(doc, s[3:].strip(), 1); cjk += count_cjk(s); i += 1; continue
        if s.startswith("### "):
            heading(doc, s[4:].strip(), 2); cjk += count_cjk(s); i += 1; continue
        if s.startswith("> "):                                   # note
            p = doc.add_paragraph(); fmt_body(p, indent=False)
            p.paragraph_format.left_indent = Pt(12)
            add_runs_bold(p, s[2:].strip(), 10.5); cjk += count_cjk(s); i += 1; continue
        if s.startswith("- "):                                   # bullet
            p = doc.add_paragraph(style="List Bullet"); fmt_body(p, indent=False)
            add_runs_bold(p, s[2:].strip(), 12); cjk += count_cjk(s); i += 1; continue
        # reference line or normal paragraph
        is_ref = bool(re.match(r"^\[\d+\]", s))
        p = doc.add_paragraph(); fmt_body(p, indent=not is_ref)
        add_runs_bold(p, s, 12, cite=not is_ref); cjk += count_cjk(s); i += 1

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[docx] wrote {out_path}")
    print(f"[docx] Chinese characters (正文+标题): {cjk:,}  ({'OK ≥3000' if cjk >= 3000 else 'BELOW 3000!'})")


def count_cjk(s):
    return sum(1 for ch in s if "一" <= ch <= "鿿")


if __name__ == "__main__":
    md = sys.argv[1] if len(sys.argv) > 1 else str(REPO / "doc/报告.md")
    out = sys.argv[2] if len(sys.argv) > 2 else str(REPO / "报告.docx")
    build(md, out)
